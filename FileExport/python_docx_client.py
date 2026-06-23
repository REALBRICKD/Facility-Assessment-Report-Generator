"""
Fetches user inputs from the Streamlit web interface and generates a facility assessment report in DOCX format using python-docx.
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

def _display_value(value, fallback="N/A"):
    if value in (None, ""):
        return fallback
    return str(value)

def _build_rows(report_data):
    claims = report_data["claims_summary"]
    return [
        ("Name of Facility", report_data.get("facility_name")),
        ("CCN", report_data.get("ccn")),
        ("State", report_data.get("state")),
        ("Location", report_data.get("location")),
        ("EMR", report_data.get("emr")),
        ("Census Capacity", report_data.get("census_capacity")),
        ("Current Census", report_data.get("current_census")),
        ("Type of Patient", report_data.get("patient_type")),
        ("Previous Coverage from MedElite", report_data.get("previous_coverage")),
        ("Previous Provider Performance from MedElite", report_data.get("previous_provider_performance")),
        ("Medical Coverage", report_data.get("medical_coverage")),
        ("Overall Star Rating", report_data.get("overall_rating")),
        ("Health Inspection", report_data.get("health_inspection_rating")),
        ("Staffing", report_data.get("staffing_rating")),
        ("Quality of Resident Care", report_data.get("quality_of_resident_care_rating")),
        ("Short Term Hospitalization", claims.get("str_hosp_score")),
        ("STR National Avg. for Hospitalization", claims.get("str_hosp_national_avg")),
        ("STR State National Avg. for Hospitalization", claims.get("str_hosp_state_avg")),
        ("Short Term ED Visit", claims.get("str_ed_score")),
        ("STR ED Visits National Avg.", claims.get("str_ed_national_avg")),
        ("STR ED Visits State Avg.", claims.get("str_ed_state_avg")),
        ("LT Hospitalization", claims.get("lt_hosp_score")),
        ("LT National Avg. for Hospitalization", claims.get("lt_hosp_national_avg")),
        ("LT State National Avg. for Hospitalization", claims.get("lt_hosp_state_avg")),
        ("ED Visit", claims.get("lt_ed_score")),
        ("LT ED Visits National Avg.", claims.get("lt_ed_national_avg")),
        ("LT ED Visits State Avg.", claims.get("lt_ed_state_avg")),
    ]

def build_docx_export(report_data):
    """
    builds document in buffer and returns it as a byte stream
    """
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("INFINITE — Managed by MEDELITE")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Helvetica"

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("FACILITY ASSESSMENT SNAPSHOT")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = "Helvetica"

    state_paragraph = document.add_paragraph()
    state_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    state_run = state_paragraph.add_run(_display_value(report_data.get("state")))
    state_run.bold = True
    state_run.font.size = Pt(12)
    state_run.font.name = "Helvetica"

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Value"

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for label, value in _build_rows(report_data):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = _display_value(value)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()